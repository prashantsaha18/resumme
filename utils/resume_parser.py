"""
utils/resume_parser.py
------------------------
Extracts plain text from uploaded resume files (PDF or DOCX) so they can
be fed into the ML models (classification, ATS scoring, job matching).
"""

import io
from typing import Union

import docx
import pdfplumber


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Also pull text from tables, since resumes sometimes use table layouts
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_resume_text(uploaded_file) -> str:
    """
    Accepts a Streamlit UploadedFile object and dispatches to the right
    extractor based on file extension.
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload a PDF, DOCX, or TXT file.")
